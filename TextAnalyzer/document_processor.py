import PyPDF2
from docx import Document
import os
import re
from typing import Tuple, Dict, Any
from fastapi import HTTPException

class DocumentProcessor:
    """
    Handles different document types: PDF, DOCX, TXT
    """
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[str, int]:
        """
        Extract text from PDF file and return text + page count
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
                
                return text, len(pdf_reader.pages)
                
        except Exception as e:
            raise HTTPException(
                status_code=500, 
                detail=f"PDF extraction failed: {str(e)}"
            )
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> Tuple[str, int]:
        """
        Extract text from DOCX file
        """
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            # Rough page estimation (500 words per page)
            word_count = len(text.split())
            estimated_pages = max(1, word_count // 500)
            
            return text, estimated_pages
            
        except Exception as e:
            raise HTTPException(
                status_code=500, 
                detail=f"DOCX extraction failed: {str(e)}"
            )
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> Tuple[str, int]:
        """
        Extract text from TXT file
        """
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Rough page estimation
            word_count = len(text.split())
            estimated_pages = max(1, word_count // 500)
            
            return text, estimated_pages
            
        except UnicodeDecodeError:
            # Try other encodings
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                word_count = len(text.split())
                estimated_pages = max(1, word_count // 500)
                
                return text, estimated_pages
                
            except Exception as e:
                raise HTTPException(
                    status_code=500, 
                    detail=f"TXT extraction failed: {str(e)}"
                )
        except Exception as e:
            raise HTTPException(
                status_code=500, 
                detail=f"TXT extraction failed: {str(e)}"
            )
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """
        Determine file type from extension
        """
        if not filename:
            return "unknown"
        
        ext = filename.lower().split('.')[-1]
        return ext
    
    @staticmethod
    def is_supported_file_type(filename: str) -> bool:
        """
        Check if file type is supported
        """
        supported_types = ['pdf', 'docx', 'txt', 'doc']
        file_type = DocumentProcessor.get_file_type(filename)
        return file_type in supported_types
    
    @staticmethod
    def extract_text(file_path: str, filename: str) -> Dict[str, Any]:
        """
        Main method to extract text from any supported document type
        """
        file_type = DocumentProcessor.get_file_type(filename)
        
        if file_type == 'pdf':
            text, pages = DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            text, pages = DocumentProcessor.extract_text_from_docx(file_path)
        elif file_type == 'txt':
            text, pages = DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {file_type}. Supported: PDF, DOCX, TXT"
            )
        
        return {
            "text": text,
            "page_count": pages,
            "file_type": file_type,
            "word_count": len(text.split()),
            "character_count": len(text)
        }